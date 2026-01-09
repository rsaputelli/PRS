# Master Gig App.py

# NOTE: This file is an admin-only internal bootstrap/orchestrator.
# Logic here is intentionally inline for stability; refactor later.

import os, io, smtplib, ssl, datetime as dt
import streamlit as st
from email.message import EmailMessage
from docx import Document
from supabase import create_client, Client
from auth_helper import restore_session, require_admin

if st.session_state.get("force_logged_out"):
    st.error("You have been logged out.")
    st.stop()


# ─────────────── Config (safe secrets + env) ───────────────
def _get_secret(name: str, default: str | None = None, required: bool = False) -> str | None:
    """Prefer st.secrets → env var → default. Stop app if required and missing."""
    val = None
    if hasattr(st, "secrets") and name in st.secrets:
        val = st.secrets.get(name)
    if val is None:
        val = os.getenv(name, default)
    if required and (val is None or str(val).strip() == ""):
        st.error(
            f"Missing required setting: `{name}`.\n\n"
            "Add it under **Settings → Secrets** in Streamlit Cloud or set it as an environment variable locally.\n"
            "Required: SUPABASE_URL, SUPABASE_ANON_KEY. Optional: SUPABASE_SERVICE_KEY, SMTP_*."
        )
        st.stop()
    return val

TEMPLATE_PATH = "PRS_Contract_Template.docx"

# Email (defaults OK)
SENDER_NAME  = "Philly Rock and Soul"
SENDER_EMAIL = _get_secret("PRS_MAIL_FROM", default="prsbandinfo@gmail.com")
SMTP_HOST    = _get_secret("SMTP_HOST", default="smtp.gmail.com")
SMTP_PORT    = int(_get_secret("SMTP_PORT", default="587"))
SMTP_USER    = _get_secret("SMTP_USER", default=SENDER_EMAIL)
SMTP_PASS    = _get_secret("SMTP_PASS")  # optional

# Supabase (REQUIRED)
SUPABASE_URL         = _get_secret("SUPABASE_URL", required=True)
SUPABASE_ANON_KEY    = _get_secret("SUPABASE_ANON_KEY", required=True)
SUPABASE_SERVICE_KEY = _get_secret("SUPABASE_SERVICE_KEY")  # optional

sb: Client = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)
sb_svc: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY) if SUPABASE_SERVICE_KEY else sb

# 🔑 Restore Supabase session on every load
user, session = restore_session()

# 🔐 Enforce admin-only access (hard stop, no UI)
user, session, user_id = require_admin()
if not user:
    st.stop()

# 🔎 SESSION DIAGNOSTIC (temporary)
# st.write("🔎 SESSION DIAGNOSTIC")
st.write("supabase_user:", st.session_state.get("supabase_user"))
# st.write("sb_access_token:", st.session_state.get("sb_access_token"))
# st.write("sb_refresh_token:", st.session_state.get("sb_refresh_token"))


# ─────────────── Helpers ───────────────
def merge_docx(template_path: str, variables: dict) -> bytes:
    doc = Document(template_path)
    repl = {f"${{{k}}}": str(v) for k, v in variables.items()}
    # paragraphs
    for p in doc.paragraphs:
        for k, v in repl.items():
            if k in p.text:
                for r in p.runs:
                    r.text = r.text.replace(k, v)
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in repl.items():
                    if k in cell.text:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = r.text.replace(k, v)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def send_email_smtp(to_addrs, subject, body_html, attachment_name, attachment_bytes, cc_addrs=None):
    msg = EmailMessage()
    msg["From"] = f"{SENDER_NAME} <{SENDER_EMAIL}>"
    msg["To"] = ", ".join(to_addrs if isinstance(to_addrs, list) else [to_addrs])
    if cc_addrs: msg["Cc"] = ", ".join(cc_addrs)
    msg["Subject"] = subject
    msg.set_content("HTML mail required.")
    msg.add_alternative(body_html, subtype="html")
    msg.add_attachment(attachment_bytes, maintype="application",
                       subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                       filename=attachment_name)
    ctx = ssl.create_default_context()
    with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as s:
        s.starttls(context=ctx); s.login(SMTP_USER, SMTP_PASS); s.send_message(msg)

# st.write("DEBUG: role_is_admin version = id-column")

def role_is_admin(user_id: str) -> bool:
    # profiles(role: 'admin' | 'standard')
    res = (
        sb.table("profiles")
        .select("role")
        .eq("id", user_id)   # ✅ correct column
        .execute()
    )

    if not res.data:
        return False

    return res.data[0]["role"] == "admin"


def load_gig_view(gig_id: str) -> dict:
    """Join gigs + venues + payments to build merge fields."""
    g = sb.table("gigs").select("*").eq("id", gig_id).single().execute().data
    if not g: raise RuntimeError("Gig not found or RLS blocked.")
    v = sb.table("venues").select("name,address,venue_type").eq("id", g["venue_id"]).single().execute().data
    pays = sb.table("gig_payments").select("*").eq("gig_id", gig_id).order("due_on").execute().data or []
    # Optional private details:
    gp = sb.table("gigs_private").select("contracted_amount").eq("gig_id", gig_id).maybe_single().execute().data

    # Merge fields (adapt names to your columns)
    # Expecting: event_date (date), time blocks, band_size, num_vocalists, overtime_rate, package fields, etc.
    # You can store these in gigs or a related packages table.
    return {
        "gig_id": gig_id,
        "event_type": g.get("event_type",""),
        "event_date": dt.date.fromisoformat(g["date_start"][:10]).strftime("%A, %B %-d, %Y") if g.get("date_start") else "",
        "venue_name": v["name"] if v else "",
        "venue_address": v["address"] if v else "",
        "venue_type": (v.get("venue_type") if v and v.get("venue_type") else (g.get("is_indoor") and "Indoor" or "Outdoor")),
        "performance_duration_hours": g.get("duration_hours",""),
        "cocktail_hours": g.get("cocktail_hours",""),
        "full_band_hours": g.get("full_band_hours",""),
        "cocktail_start": g.get("cocktail_start",""),
        "cocktail_end": g.get("cocktail_end",""),
        "reception_start": g.get("reception_start",""),
        "reception_end": g.get("reception_end",""),

        "contract_total": gp.get("contracted_amount") if gp else g.get("fee_total",""),
        "credit_card_fee_pct": g.get("credit_card_fee_pct","4.5"),

        "package_name": g.get("package_name",""),
        "package_price": g.get("package_price",""),
        "band_size": g.get("band_size",""),
        "num_vocalists": g.get("num_vocalists",""),
        "overtime_rate": g.get("overtime_rate","$500"),
        "stage_space": g.get("stage_space","10 × 20 ft"),
        "access_lead_hours": g.get("access_lead_hours","3"),

        # Map up to 3 deposits + final (if your UI collects arbitrary N, keep first 3 + final for template)
        "deposit1_amt": pays[0]["amount"] if len(pays) > 0 and pays[0].get("type") in ("Deposit","deposit") else "",
        "deposit1_due": (dt.date.fromisoformat(pays[0]["due_on"]).strftime("%B %-d, %Y") if len(pays)>0 and pays[0].get("due_on") else "At signing"),
        "deposit2_amt": pays[1]["amount"] if len(pays) > 1 and pays[1].get("type") in ("Deposit","deposit") else "",
        "deposit2_due": (dt.date.fromisoformat(pays[1]["due_on"]).strftime("%B %-d, %Y") if len(pays)>1 and pays[1].get("due_on") else ""),
        "deposit3_amt": pays[2]["amount"] if len(pays) > 2 and pays[2].get("type") in ("Deposit","deposit") else "",
        "deposit3_due": (dt.date.fromisoformat(pays[2]["due_on"]).strftime("%B %-d, %Y") if len(pays)>2 and pays[2].get("due_on") else ""),
        "final_payment_amt": next((p["amount"] for p in pays if p.get("type","").lower()=="final"), ""),
        "final_payment_due": next((dt.date.fromisoformat(p["due_on"]).strftime("%B %-d, %Y")
                                   for p in pays if p.get("type","").lower()=="final" and p.get("due_on")), ""),

        # hosts (store on gigs or private_event_details)
        "host_names": g.get("host_names",""),
        "host1_name": g.get("host1_name",""),
        "host2_name": g.get("host2_name",""),
        "host_email": g.get("host_email",""),

        "contract_status": g.get("contract_status","Draft"),
        "created_by_user": g.get("created_by",""),
        "created_at": g.get("created_at",""),
        "contract_signed_date_band": "",
        "host_signed_date": "",
    }

def upload_contract_to_storage(gig_id: str, file_name: str, data: bytes) -> str:
    # Ensure bucket 'contracts' exists; path per gig
    path = f"{gig_id}/{file_name}"
    sb_svc.storage.from_("contracts").upload(path, data, {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "upsert": "true"})
    # If bucket is public:
    public = sb_svc.storage.from_("contracts").get_public_url(path)
    return public["publicUrl"] if isinstance(public, dict) else public

def mark_contract_sent(gig_id: str, url: str):
    sb_svc.table("gigs").update({"contract_status": "Sent", "contract_url": url}).eq("id", gig_id).execute()

def upsert_package(name: str, default_price: str | float):
    # packages(name unique, is_active, default_price)
    sb_svc.table("packages").upsert({
        "name": name.strip(),
        "is_active": True,
        "default_price": default_price
    }, on_conflict="name").execute()

# ─────────────── UI ───────────────
st.set_page_config(page_title="Master App (bootstrap use only, no user interaction)", page_icon="🎶", layout="wide")
st.title("Master Gig App (Admin)")

# Auth/role check (replace with your auth session)
user_id = st.session_state.get("user_id")


gig_id = st.query_params.get("gig_id", [None])[0]

if not gig_id:
    st.info("No gig selected. Open this page from a gig context.")
    st.stop()

try:
    data = load_gig_view(gig_id)
except Exception as e:
    st.error(f"Could not load gig: {e}")
    st.stop()

# Show summary + admin overrides
colL, colR = st.columns([2,1])
with colL:
    st.subheader("Contract Summary")
    st.write({
        "Event": f'{data["event_type"]} on {data["event_date"]}',
        "Venue": f'{data["venue_name"]}, {data["venue_address"]} ({data["venue_type"]})',
        "Schedule": {
            "Cocktail": f'{data["cocktail_start"]} – {data["cocktail_end"]}',
            "Reception": f'{data["reception_start"]} – {data["reception_end"]}',
            "Duration (hrs)": data["performance_duration_hours"],
        },
        "Package": f'{data["package_name"]} ({data["package_price"]})',
        "Band": f'{data["band_size"]}-piece, {data["num_vocalists"]} vocalist(s)',
        "Overtime": data["overtime_rate"] + "/hour",
        "Total": data["contract_total"],
        "Deposits": [
            (data["deposit1_amt"], data["deposit1_due"]),
            (data["deposit2_amt"], data["deposit2_due"]),
            (data["deposit3_amt"], data["deposit3_due"]),
        ],
        "Final Payment": (data["final_payment_amt"], data["final_payment_due"]),
        "Hosts": data["host_names"],
        "Emails": data["host_email"],
    })

    with st.expander("Optional Admin Overrides (+ OTHER pattern)"):
        # Package dropdown from DB
        pkg_rows = sb.table("packages").select("name, default_price").eq("is_active", True).order("name").execute().data or []
        pkg_names = [r["name"] for r in pkg_rows] + ["OTHER"]
        sel = st.selectbox("Package", pkg_names, index=(pkg_names.index(data["package_name"]) if data.get("package_name") in pkg_names else len(pkg_names)-1))
        if sel == "OTHER":
            new_pkg = st.text_input("New Package Name", value=data.get("package_name",""))
            new_price = st.text_input("New Package Default Price", value=str(data.get("package_price","")))
            if st.button("Save NEW Package to DB"):
                if new_pkg.strip():
                    upsert_package(new_pkg, new_price or 0)
                    st.success("Saved. Reload to see in dropdown.")
                else:
                    st.warning("Package name required.")
            # use current entry for this document
            data["package_name"] = new_pkg or data["package_name"]
            data["package_price"] = new_price or data["package_price"]
        else:
            data["package_name"] = sel
            match = next((r for r in pkg_rows if r["name"] == sel), None)
            data["package_price"] = match["default_price"] if match else data["package_price"]

        # Editable totals
        data["contract_total"] = st.text_input("Contract Total", str(data["contract_total"]))
        data["venue_type"] = st.selectbox("Indoor/Outdoor", ["Indoor","Outdoor","Mixed"], index=["Indoor","Outdoor","Mixed"].index(data.get("venue_type","Indoor")))
        data["overtime_rate"] = st.text_input("Overtime Rate", str(data["overtime_rate"]))

with colR:
    st.subheader("Generate Draft")
    if not os.path.exists(TEMPLATE_PATH):
        st.error("Template file missing. Upload PRS_Contract_Template.docx to app root.")
        st.stop()
    draft_bytes = merge_docx(TEMPLATE_PATH, data)
    st.download_button("⬇️ Download Draft DOCX",
                       draft_bytes,
                       file_name=f"PRS_Contract_{gig_id}.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       use_container_width=True)

st.markdown("---")
st.subheader("Admin Review")
ok_1 = st.checkbox("Verified dates/times/venue and Indoor/Outdoor.")
ok_2 = st.checkbox("Verified package, price, total, deposits, final payment.")
ok_3 = st.checkbox("Verified host names and recipient email(s).")
ready = ok_1 and ok_2 and ok_3

col1, col2, col3 = st.columns([1,1,2])
with col1:
    store_now = st.button("💾 Save to Storage (Draft)")
with col2:
    send_now = st.button("✅ Approve & Send", disabled=not ready)

storage_url = None
if store_now:
    storage_url = upload_contract_to_storage(gig_id, f"PRS_Contract_{gig_id}.docx", draft_bytes)
    st.success(f"Saved draft to storage:\n{storage_url}")

if send_now and ready:
    storage_url = upload_contract_to_storage(gig_id, f"PRS_Contract_{gig_id}.docx", draft_bytes)
    mark_contract_sent(gig_id, storage_url)
    to_emails = [e.strip() for e in (data.get("host_email") or "").split(",") if e.strip()]
    cc_emails = [SENDER_EMAIL]
    subj = f"Philly Rock and Soul – Contract for {data['event_type']} on {data['event_date']}"
    html = f"""
    <p>Hi {data['host_names']},</p>
    <p>Attached is your contract for <b>{data['event_type']}</b> on <b>{data['event_date']}</b> at <b>{data['venue_name']}</b>.</p>
    <p>Please reply to confirm; we’ll countersign. Thank you!</p>
    <p>— Philly Rock and Soul</p>
    <p><small>Download link: <a href="{storage_url}">{storage_url}</a></small></p>
    """
    try:
        send_email_smtp(to_emails, subj, html, f"PRS_Contract_{gig_id}.docx", draft_bytes, cc_addrs=cc_emails)
        st.success("Approved and sent. A copy was CC’d internally.")
    except Exception as e:
        st.error(f"Email failed: {e}")