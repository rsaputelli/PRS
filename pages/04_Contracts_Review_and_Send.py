import os
import io
import streamlit as st
from docx import Document
from docx.shared import Inches
from supabase import create_client, Client
from tools.contracts import build_private_contract_context

# Build Supabase client exactly like other pages
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_ANON_KEY = st.secrets["SUPABASE_ANON_KEY"]
sb: Client = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)

ASSETS_DIR = os.path.join(os.path.dirname(__file__), "..", "assets")
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "assets", "PRS_Contract_Template.docx")


# ---------------------------------------------------------------------
# Replace {{tokens}} inside a paragraph
# Handles text + signature + logo
# ---------------------------------------------------------------------
def _replace_tokens(paragraph, ctx):
    for key, value in ctx.items():
        token = f"{{{{{key}}}}}"

        if token in paragraph.text:

            # Clear existing text runs in paragraph
            for run in paragraph.runs:
                run.text = ""

            # Special case: signature image
            if key == "signature":
                sig_path = ctx.get("signature_image_path")
                if sig_path and os.path.exists(sig_path):
                    try:
                        run = paragraph.add_run()
                        run.add_picture(sig_path, width=Inches(1.75))
                        return
                    except Exception:
                        paragraph.add_run("[Signature Missing]")
                        return

            # Special case: logo
            if key == "logo":
                logo_path = ctx.get("logo_image_path")
                if logo_path and os.path.exists(logo_path):
                    try:
                        run = paragraph.add_run()
                        run.add_picture(logo_path, width=Inches(2.5))
                        return
                    except Exception:
                        paragraph.add_run("[Logo Missing]")
                        return

            # Normal text replacement
            replacement = str(value) if value is not None else ""
            paragraph.add_run(replacement)


# ---------------------------------------------------------------------
# Merge DOCX Template
# ---------------------------------------------------------------------
def merge_docx_template(ctx):
    doc = Document(TEMPLATE_PATH)

    # Loop paragraphs
    for paragraph in doc.paragraphs:
        _replace_tokens(paragraph, ctx)

    # Loop tables too (many DOCX templates hide tokens in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_tokens(paragraph, ctx)

    return doc


# ---------------------------------------------------------------------
# Streamlit Page
# ---------------------------------------------------------------------
def main():
    st.title("📤 Contract Review & Send")

    if "user" not in st.session_state or not st.session_state["user"]:
        st.error("Please sign in.")
        st.stop()

    sb = get_sb_client()

    st.write("Select a gig to generate its contract:")

    # -------------------------------------------------------------
    # Gig selection dropdown
    # -------------------------------------------------------------
    gigs_resp = sb.table("gigs").select("id, title, event_date").order("event_date").execute()
    gigs = gigs_resp.data if gigs_resp.data else []

    gig_options = {f"{g['event_date']} — {g['title']}": g["id"] for g in gigs}
    selected_label = st.selectbox("Choose Event", list(gig_options.keys()) if gig_options else [])

    if not selected_label:
        st.info("Pick an event to continue.")
        st.stop()

    gig_id = gig_options[selected_label]

    # -------------------------------------------------------------
    # Build context (from tools/contracts.py)
    # -------------------------------------------------------------
    ctx = build_private_contract_context(sb, gig_id)

    # Add logo/sig fields expected by the templater
    ctx["signature"] = "signature"
    ctx["logo"] = "logo"
    ctx["logo_image_path"] = os.path.join(ASSETS_DIR, "prs_logo.png")

    st.subheader("📄 Contract Preview (Text Only)")
    st.markdown("---")

    # Render preview using your renderer
    from tools.contracts import _render_contract_preview
    st.markdown(_render_contract_preview(ctx))

    with st.expander("Show merged context (debug)", expanded=False):
        st.json(ctx)

    st.markdown("---")

    # -------------------------------------------------------------
    # Generate DOCX
    # -------------------------------------------------------------
    if st.button("Generate Contract DOCX"):
        doc = merge_docx_template(ctx)

        # Save to in-memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("Contract generated successfully!")

        st.download_button(
            label="⬇️ Download Contract",
            data=buffer,
            file_name=f"Contract_{gig_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        st.info("Next step: we'll wire up the email-to-client workflow.")

if __name__ == "__main__":
    main()

